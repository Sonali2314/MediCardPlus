"""
Python NLP Service for Medical Data Extraction
Extracts medical information from PDF documents
"""

import sys
import json
import re
from pathlib import Path

# Try to import pdf and docx libraries
try:
    import PyPDF2
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    sys.stderr.write("Warning: PyPDF2 not installed - PDF extraction disabled\n")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    sys.stderr.write("Warning: python-docx not installed - DOCX extraction disabled\n")

# Medical keywords dictionary
MEDICAL_KEYWORDS = {
    "symptoms": [
        "fever", "cough", "headache", "pain", "nausea", "vomiting", "dizziness",
        "fatigue", "weakness", "shortness of breath", "chest pain", "abdominal pain",
        "diarrhea", "constipation", "rash", "itching", "swelling", "inflammation",
        "chills", "sweating", "anxiety", "depression", "insomnia", "tremor",
        "numbness", "tingling", "palpitations", "jaundice", "breathlessness",
        "loss of appetite", "weight loss", "chills", "night sweats", "cough"
    ],
    "medications": [
        "aspirin", "ibuprofen", "paracetamol", "acetaminophen", "amoxicillin",
        "penicillin", "metformin", "lisinopril", "atorvastatin", "metoprolol",
        "omeprazole", "ranitidine", "loratadine", "cetirizine", "diphenhydramine",
        "insulin", "methotrexate", "azithromycin", "ciprofloxacin", "doxycycline",
        "prednisone", "levothyroxine", "amlodipine", "warfarin", "clopidogrel"
    ],
    "diagnoses": [
        "diabetes", "hypertension", "asthma", "arthritis", "pneumonia", "influenza",
        "covid", "bronchitis", "gastritis", "ulcer", "migraine", "anemia",
        "hypothyroidism", "hyperthyroidism", "kidney disease", "heart disease",
        "stroke", "cancer", "tuberculosis", "malaria", "dengue", "typhoid",
        "hepatitis", "cirrhosis", "appendicitis", "gallstones", "obesity"
    ],
    "tests": [
        "blood test", "x-ray", "ct scan", "mri", "ultrasound", "ecg", "ekg",
        "biopsy", "endoscopy", "colonoscopy", "hemoglobin", "glucose", "cholesterol",
        "triglyceride", "liver function", "kidney function", "urinalysis",
        "tsh", "t3", "t4", "white blood cell", "red blood cell", "platelet count",
        "hba1c", "serum creatinine", "uric acid", "vitamin d", "vitamin b12"
    ]
}


def extract_text_from_pdf(file_path):
    """Extract text from PDF file using PyPDF2"""
    try:
        if not HAS_PYPDF:
            sys.stderr.write("Warning: PyPDF2 not installed\n")
            return ""
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            sys.stderr.write(f"PDF has {num_pages} pages\n")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    sys.stderr.write(f"Error extracting page {page_num}: {str(e)}\n")
        
        sys.stderr.write(f"Extracted {len(text)} characters total\n")
        return text
    except Exception as e:
        sys.stderr.write(f"PDF extraction error: {str(e)}\n")
        return ""


def extract_text_from_docx(file_path):
    """Extract text from DOCX file using python-docx"""
    try:
        if not HAS_DOCX:
            sys.stderr.write("Warning: python-docx not installed\n")
            return ""
        
        text = ""
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
            text += "\n"
        
        sys.stderr.write(f"Extracted {len(text)} characters from DOCX\n")
        return text
    except Exception as e:
        sys.stderr.write(f"DOCX extraction error: {str(e)}\n")
        return ""


def extract_medical_info(text):
    """Extract medical keywords from text using regex matching"""
    if not text:
        return {
            "symptoms": [],
            "medications": [],
            "diagnoses": [],
            "tests": [],
            "raw_text_length": 0
        }
    
    text_lower = text.lower()
    
    extracted = {
        "symptoms": [],
        "medications": [],
        "diagnoses": [],
        "tests": [],
        "raw_text_length": len(text)
    }
    
    # Extract symptoms
    for keyword in MEDICAL_KEYWORDS["symptoms"]:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, text_lower):
            if keyword not in extracted["symptoms"]:
                extracted["symptoms"].append(keyword)
    
    # Extract medications
    for keyword in MEDICAL_KEYWORDS["medications"]:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, text_lower):
            if keyword not in extracted["medications"]:
                extracted["medications"].append(keyword)
    
    # Extract diagnoses
    for keyword in MEDICAL_KEYWORDS["diagnoses"]:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, text_lower):
            if keyword not in extracted["diagnoses"]:
                extracted["diagnoses"].append(keyword)
    
    # Extract tests
    for keyword in MEDICAL_KEYWORDS["tests"]:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, text_lower):
            if keyword not in extracted["tests"]:
                extracted["tests"].append(keyword)
    
    return extracted


def process_file(file_path):
    """Main function to process a medical file"""
    
    file_path = Path(file_path)
    
    if not file_path.exists():
        return {
            "success": False,
            "error": f"File not found: {file_path}"
        }
    
    try:
        # Extract text based on file type
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            sys.stderr.write(f"Processing PDF: {file_path}\n")
            text = extract_text_from_pdf(str(file_path))
        elif file_ext in ['.docx', '.doc']:
            sys.stderr.write(f"Processing DOCX: {file_path}\n")
            text = extract_text_from_docx(str(file_path))
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_ext}"
            }
        
        if not text or len(text.strip()) == 0:
            sys.stderr.write("No text extracted from file\n")
            return {
                "success": True,
                "extracted_data": {
                    "symptoms": [],
                    "medications": [],
                    "diagnoses": [],
                    "tests": [],
                    "raw_text_length": 0
                },
                "warning": f"No text could be extracted from the {file_ext} file"
            }
        
        # Extract medical information
        medical_info = extract_medical_info(text)
        
        sys.stderr.write(f"Extracted: {len(medical_info['symptoms'])} symptoms, {len(medical_info['medications'])} medications, {len(medical_info['diagnoses'])} diagnoses, {len(medical_info['tests'])} tests\n")
        
        return {
            "success": True,
            "extracted_data": medical_info
        }
    
    except Exception as e:
        sys.stderr.write(f"Error processing file: {str(e)}\n")
        return {
            "success": False,
            "error": str(e)
        }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "No file path provided"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = process_file(file_path)
    print(json.dumps(result))

